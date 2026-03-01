"""
VoltLegal — Gemini Service
PDF, Image, Word doc, and Voice analysis using Google Gemini 2.5 Flash.
"""

import os
import io
import time
import logging
import google.generativeai as genai
from dotenv import load_dotenv
from PIL import Image

load_dotenv()

GEMINI_API_KEY = os.getenv("GEMINI_API_KEY")
logger = logging.getLogger(__name__)

if GEMINI_API_KEY:
    genai.configure(api_key=GEMINI_API_KEY)

GEMINI_MODEL = "gemini-2.5-flash"


def _get_model():
    """Get a configured Gemini model instance."""
    if not GEMINI_API_KEY:
        raise RuntimeError("GEMINI_API_KEY not found in .env")
    return genai.GenerativeModel(GEMINI_MODEL)


def _call_gemini(contents, temperature=0.3, max_tokens=4096, retries=3):
    """Call Gemini API with automatic retry and exponential backoff."""
    model = _get_model()
    last_err = None
    for attempt in range(retries):
        try:
            response = model.generate_content(
                contents,
                generation_config=genai.GenerationConfig(
                    temperature=temperature,
                    max_output_tokens=max_tokens,
                ),
            )
            return response.text
        except Exception as e:
            last_err = e
            wait = 2 ** attempt
            logger.warning(f"Gemini API attempt {attempt+1} failed: {e}. Retrying in {wait}s...")
            time.sleep(wait)
    raise RuntimeError(f"Gemini API failed after {retries} attempts: {last_err}")


# ─── Prompts ─────────────────────────────────────────────────────────────────

PDF_ANALYSIS_PROMPT = """You are VoltLegal, an Indian legal document analyzer.

Analyze this legal document and provide:

📋 **Document Type:**
[What kind of document is this? E.g., Rent Agreement, Court Notice, FIR, Contract, etc.]

📝 **Summary:**
[Plain-language summary of what this document says. Explain every important point.]

⚠️ **Key Clauses & Red Flags:**
[Highlight any important clauses, obligations, deadlines, or potentially unfair terms]

📜 **Relevant Laws:**
• [Mention which Indian laws/acts govern this type of document]

✅ **What You Should Know:**
[Practical advice — what to watch out for, what rights the person has]

STRICT RULES:
- Explain everything in simple language as if talking to someone with no legal knowledge.
- If the document appears to involve illegal activity, flag it clearly.
- Never encourage illegal activities.
- Provide legal INFORMATION, not legal ADVICE.
- Focus on Indian law unless the document clearly pertains to another jurisdiction."""

WORD_DOC_ANALYSIS_PROMPT = """You are VoltLegal, an Indian legal document analyzer.

Analyze this legal document content and provide:

📋 **Document Type:**
[What kind of document is this? E.g., Rent Agreement, Contract, Legal Notice, etc.]

📝 **Summary:**
[Plain-language summary of what this document says. Explain every important point.]

⚠️ **Key Clauses & Red Flags:**
[Highlight any important clauses, obligations, deadlines, or potentially unfair terms]

📜 **Relevant Laws:**
• [Mention which Indian laws/acts govern this type of document]

✅ **What You Should Know:**
[Practical advice — what to watch out for, what rights the person has]

STRICT RULES:
- Explain everything in simple language as if talking to someone with no legal knowledge.
- If the document appears to involve illegal activity, flag it clearly.
- Never encourage illegal activities.
- Provide legal INFORMATION, not legal ADVICE.
- Focus on Indian law unless the document clearly pertains to another jurisdiction."""

IMAGE_ANALYSIS_PROMPT = """You are VoltLegal, an Indian legal document analyzer.

Look at this image carefully. It may be a photo of a legal document, court notice, FIR, agreement, or any legal paper.

Do your best to read and analyze ALL the text visible in the image, even if some parts are slightly unclear. Try to identify:

📋 **Document Type:**
[What kind of document does this appear to be?]

📝 **Content Summary:**
[Read and explain everything written in this document in simple language. Extract as much text as possible.]

📜 **Relevant Laws:**
• [Which Indian laws/acts are relevant to this document]

✅ **What Should You Do:**
[Practical steps the person should consider]

STRICT RULES:
- ALWAYS attempt to read and analyze the image, even if parts are unclear or the quality is not perfect.
- If some parts are hard to read, still analyze what you CAN read and note which parts were unclear.
- NEVER refuse to analyze an image just because it's not perfectly clear. Always try your best.
- Explain in simple language.
- Never encourage illegal activities.
- Provide legal INFORMATION, not legal ADVICE.
- If this does not appear to be a legal document, still describe what you see in the image."""

VOICE_TRANSCRIPTION_PROMPT = """Listen to this audio message carefully and transcribe it accurately.
The user is speaking about a legal matter or question.
Return ONLY the transcribed text, nothing else. Do not add commentary.
If the speech is in Hindi, Telugu, or another Indian language, transcribe it in that language using the appropriate script.
If you cannot understand certain parts, mark them as [unclear]."""


# ─── Service Functions ───────────────────────────────────────────────────────

def analyze_pdf(pdf_bytes: bytes, filename: str = "document.pdf") -> str:
    """Analyze a PDF document using Gemini."""
    pdf_file = genai.upload_file(
        io.BytesIO(pdf_bytes),
        mime_type="application/pdf",
        display_name=filename,
    )

    try:
        result = _call_gemini([PDF_ANALYSIS_PROMPT, pdf_file])
    finally:
        try:
            genai.delete_file(pdf_file.name)
        except Exception:
            pass

    return result


def analyze_image(image_bytes: bytes) -> str:
    """Analyze a photo of a legal document using Gemini Vision."""
    # Detect mime type from image bytes
    img = Image.open(io.BytesIO(image_bytes))
    fmt = img.format or "JPEG"
    mime_map = {"JPEG": "image/jpeg", "PNG": "image/png", "GIF": "image/gif", "WEBP": "image/webp"}
    mime_type = mime_map.get(fmt.upper(), "image/jpeg")

    # Convert to JPEG if format is unusual
    if fmt.upper() not in mime_map:
        buf = io.BytesIO()
        img.convert("RGB").save(buf, format="JPEG")
        image_bytes = buf.getvalue()
        mime_type = "image/jpeg"

    # Upload as file for better analysis
    uploaded_file = genai.upload_file(
        io.BytesIO(image_bytes),
        mime_type=mime_type,
        display_name="legal_document_image",
    )

    try:
        result = _call_gemini([IMAGE_ANALYSIS_PROMPT, uploaded_file])
    finally:
        try:
            genai.delete_file(uploaded_file.name)
        except Exception:
            pass

    return result


def analyze_word_doc(doc_bytes: bytes, filename: str = "document.docx") -> str:
    """Analyze a Word document (.docx) by extracting text and sending to Gemini."""
    # Extract text from docx
    try:
        from docx import Document
        doc = Document(io.BytesIO(doc_bytes))
        full_text = []
        for para in doc.paragraphs:
            if para.text.strip():
                full_text.append(para.text.strip())

        # Also extract text from tables
        for table in doc.tables:
            for row in table.rows:
                row_text = [cell.text.strip() for cell in row.cells if cell.text.strip()]
                if row_text:
                    full_text.append(" | ".join(row_text))

        extracted_text = "\n\n".join(full_text)
    except Exception as e:
        raise RuntimeError(f"Could not read Word document: {e}")

    if not extracted_text.strip():
        return ("The document appears to be empty or contains only images. "
                "Please try uploading as a PDF or send photos of the pages.")

    return _call_gemini([
        WORD_DOC_ANALYSIS_PROMPT,
        f"Document filename: {filename}\n\nDocument content:\n{extracted_text}"
    ])


def transcribe_voice(audio_bytes: bytes, mime_type: str = "audio/ogg") -> str:
    """Transcribe a voice message using Gemini."""
    uploaded_file = genai.upload_file(
        io.BytesIO(audio_bytes),
        mime_type=mime_type,
        display_name="voice_message",
    )

    try:
        result = _call_gemini(
            [VOICE_TRANSCRIPTION_PROMPT, uploaded_file],
            temperature=0.1,
            max_tokens=2048,
        )
    finally:
        try:
            genai.delete_file(uploaded_file.name)
        except Exception:
            pass

    return result
